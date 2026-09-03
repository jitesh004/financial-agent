"""Generate a coherent set of synthetic statements for development and tests.

This is not a random data dump. It simulates one person's finances for 12 months
so that the pipeline can be verified end to end:

  - every account's balances actually tie out (opening + credits - debits == closing)
  - salary lands in the savings account and is then moved outward
  - the credit-card bill paid from savings appears in BOTH statements, so
    transfer detection has something real to catch
  - the home loan EMI amortizes properly, splitting into interest and principal
  - each account is rendered in a DIFFERENT file format, exercising all extractors

Run:  python backend/tools/generate_samples.py
"""

from __future__ import annotations

import csv
import random
import sys
from dataclasses import dataclass, field
from datetime import date, timedelta
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "data" / "samples"

random.seed(20260826)  # deterministic fixtures

MONTHS = 12
START = date(2025, 9, 1)

SALARY_GROSS = Decimal("225000")
SALARY_NET = Decimal("168400")


def money(value) -> Decimal:
    return Decimal(value).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


def add_months(d: date, n: int) -> date:
    y, m = divmod(d.month - 1 + n, 12)
    return date(d.year + y, m + 1, min(d.day, 28))


@dataclass
class Txn:
    when: date
    desc: str
    debit: Decimal | None = None
    credit: Decimal | None = None
    balance: Decimal = Decimal("0")
    ref: str = ""


@dataclass
class Ledger:
    name: str
    opening: Decimal
    #: Liability accounts (cards, loans) run the opposite way: a debit column
    #: entry increases what you owe. Getting this backwards is the classic
    #: statement-parsing bug, so the fixtures deliberately contain both.
    liability: bool = False
    txns: list[Txn] = field(default_factory=list)

    def add(self, when: date, desc: str, debit=None, credit=None, ref: str = "") -> Txn:
        t = Txn(when, desc, money(debit) if debit else None,
                money(credit) if credit else None, Decimal("0"), ref)
        self.txns.append(t)
        return t

    def finalize(self) -> Decimal:
        """Sort chronologically and stamp the running balance. This is what
        makes the reconciliation gate meaningful."""
        self.txns.sort(key=lambda t: (t.when, t.desc))
        bal = self.opening
        for t in self.txns:
            debit = t.debit or Decimal("0")
            credit = t.credit or Decimal("0")
            bal += (debit - credit) if self.liability else (credit - debit)
            t.balance = money(bal)
        return money(bal)


def ref(n: int) -> str:
    return f"{random.randint(10**11, 10**12 - 1)}"


# --------------------------------------------------------------------------
# Simulation
# --------------------------------------------------------------------------

MERCHANTS = {
    "groceries": ["BIGBASKET", "DMART", "ZEPTO", "BLINKIT", "RELIANCE FRESH"],
    "dining": ["SWIGGY", "ZOMATO", "STARBUCKS", "THIRD WAVE COFFEE", "TRUFFLES"],
    "transport": ["UBER INDIA", "OLA CABS", "RAPIDO", "NAMMA METRO"],
    "shopping": ["AMAZON IN", "MYNTRA", "FLIPKART", "DECATHLON", "IKEA"],
    "entertainment": ["BOOKMYSHOW", "PVR CINEMAS", "SPOTIFY"],
    "fuel": ["INDIAN OIL", "HP PETROL PUMP", "SHELL"],
    "healthcare": ["APOLLO PHARMACY", "PRACTO", "1MG"],
}

SUBSCRIPTIONS = [
    ("NETFLIX SUBSCRIPTION", Decimal("649"), 3),
    ("SPOTIFY PREMIUM", Decimal("119"), 5),
    ("AMAZON PRIME", Decimal("299"), 12),
    ("ICLOUD STORAGE", Decimal("219"), 8),
    ("CULT FIT MEMBERSHIP", Decimal("1499"), 15),
]


def simulate():
    savings = Ledger("savings", Decimal("342180.50"))
    card = Ledger("card", Decimal("0"), liability=True)
    home = Ledger("home_loan", Decimal("4185000"), liability=True)
    personal = Ledger("personal_loan", Decimal("480000"), liability=True)
    mf = Ledger("investment", Decimal("0"))

    home_rate = Decimal("0.0875") / 12
    home_emi = Decimal("38420")
    home_balance = home.opening

    pers_rate = Decimal("0.1425") / 12
    pers_emi = Decimal("16250")
    pers_balance = personal.opening

    mf_units = Decimal("0")
    nav = Decimal("48.2130")

    card_prev_due = Decimal("0")

    for m in range(MONTHS):
        month_start = add_months(START, m)

        # ---- Salary -----------------------------------------------------
        pay_day = month_start.replace(day=1) + timedelta(days=0)
        savings.add(pay_day, f"NEFT-CR-HDFC0000521-CUBYTS TECHNOLOGIES PVT LTD-SALARY {month_start:%b%y}",
                    credit=SALARY_NET, ref=ref(m))

        # ---- Home loan EMI (amortizing) --------------------------------
        interest = money(home_balance * home_rate)
        principal = money(home_emi - interest)
        home_balance = money(home_balance - principal)
        emi_day = month_start.replace(day=5)
        savings.add(emi_day, "ACH-D- HDFC LTD HOME LOAN EMI-HL4471929", debit=home_emi, ref=ref(m))
        home.add(emi_day, f"INTEREST CHARGED @ 8.75% FOR {month_start:%b %Y}", debit=interest)
        home.add(emi_day, f"EMI RECEIVED - INSTALMENT {m + 1} (PRINCIPAL {principal:,.2f})",
                 credit=home_emi, ref=ref(m))

        # ---- Personal loan EMI -----------------------------------------
        p_int = money(pers_balance * pers_rate)
        p_prin = money(pers_emi - p_int)
        pers_balance = money(pers_balance - p_prin)
        p_day = month_start.replace(day=7)
        savings.add(p_day, "ACH-D- BAJAJ FINSERV PERSONAL LOAN PL8823641", debit=pers_emi, ref=ref(m))
        personal.add(p_day, f"INTEREST CHARGED @ 14.25% FOR {month_start:%b %Y}", debit=p_int)
        personal.add(p_day, f"EMI RECEIVED - INSTALMENT {m + 1} (PRINCIPAL {p_prin:,.2f})",
                     credit=pers_emi, ref=ref(m))

        # ---- Rent -------------------------------------------------------
        savings.add(month_start.replace(day=3),
                    "IMPS-P2A-RAMESH KUMAR-HOUSE RENT", debit=Decimal("52000"), ref=ref(m))

        # ---- SIP investments (savings -> mutual fund) -------------------
        sip_day = month_start.replace(day=10)
        for fund, amt in [("PARAG PARIKH FLEXI CAP", Decimal("25000")),
                          ("UTI NIFTY 50 INDEX FUND", Decimal("15000"))]:
            savings.add(sip_day, f"ACH-D- BSE LTD SIP {fund}", debit=amt, ref=ref(m))
            nav_m = money(nav * (1 + Decimal(random.uniform(-0.03, 0.045))))
            units = (amt / nav_m).quantize(Decimal("0.0001"))
            mf_units += units
            mf.add(sip_day, f"SIP PURCHASE - {fund}", debit=amt, ref=f"{units} units @ {nav_m}")
        nav = money(nav * (1 + Decimal(random.uniform(-0.02, 0.04))))

        # ---- Credit card bill payment (appears in BOTH ledgers) ---------
        if card_prev_due > 0:
            bill_day = month_start.replace(day=18)
            savings.add(bill_day, "IMPS-P2A-ICICI BANK CREDIT CARD PAYMENT-XX4471",
                        debit=card_prev_due, ref=ref(m))
            card.add(bill_day, "PAYMENT RECEIVED - THANK YOU", credit=card_prev_due, ref=ref(m))
            card_prev_due = Decimal("0")

        # ---- Utilities ---------------------------------------------------
        for day, label, lo, hi in [
            (12, "BESCOM ELECTRICITY BILL", 2200, 4800),
            (14, "ACT FIBERNET BROADBAND", 1180, 1180),
            (16, "AIRTEL POSTPAID BILL", 899, 1499),
            (20, "BWSSB WATER CHARGES", 380, 720),
        ]:
            savings.add(month_start.replace(day=day), f"BIL/ONL/{label}",
                        debit=Decimal(random.randint(lo, hi)), ref=ref(m))

        # ---- Everyday spend on the credit card --------------------------
        month_card_spend = Decimal("0")
        for _ in range(random.randint(22, 34)):
            cat = random.choice(list(MERCHANTS.keys()))
            merchant = random.choice(MERCHANTS[cat])
            base = {
                "groceries": (450, 3200), "dining": (180, 2400), "transport": (60, 720),
                "shopping": (399, 12000), "entertainment": (250, 1800),
                "fuel": (1500, 4500), "healthcare": (150, 2800),
            }[cat]
            amt = Decimal(random.randint(*base))
            day = random.randint(1, 27)
            card.add(month_start.replace(day=1) + timedelta(days=day - 1),
                     f"{merchant}              BANGALORE IN", debit=amt)
            month_card_spend += amt

        # ---- Subscriptions on the card ----------------------------------
        for label, amt, day in SUBSCRIPTIONS:
            card.add(month_start.replace(day=day), f"{label}   MUMBAI IN", debit=amt)
            month_card_spend += amt

        # ---- Some spend direct from savings -----------------------------
        for _ in range(random.randint(3, 6)):
            merchant = random.choice(MERCHANTS[random.choice(list(MERCHANTS.keys()))])
            savings.add(month_start.replace(day=random.randint(2, 27)),
                        f"UPI/{merchant}/{ref(m)}/PAYMENT",
                        debit=Decimal(random.randint(200, 4500)), ref=ref(m))

        savings.add(month_start.replace(day=random.randint(5, 25)),
                    "ATW-CASH WITHDRAWAL-ATM ID S1BG4471",
                    debit=Decimal(random.choice([2000, 5000, 10000])), ref=ref(m))

        # ---- Quarterly savings interest ---------------------------------
        if m % 3 == 2:
            savings.add(month_start.replace(day=27), "CREDIT INTEREST CAPITALISED",
                        credit=Decimal(random.randint(1800, 3400)))

        # ---- Occasional income ------------------------------------------
        if m in (2, 7):
            savings.add(month_start.replace(day=22),
                        "NEFT-CR-FREELANCE CONSULTING INVOICE", credit=Decimal("85000"), ref=ref(m))
        if m == 9:
            savings.add(month_start.replace(day=15),
                        "NEFT-CR-CUBYTS TECHNOLOGIES-PERFORMANCE BONUS",
                        credit=Decimal("340000"), ref=ref(m))

        # Card fees, then roll the statement
        if m % 6 == 0:
            card.add(month_start.replace(day=25), "ANNUAL FEE - GST INCLUSIVE", debit=Decimal("5900"))
            month_card_spend += Decimal("5900")
        card_prev_due = money(month_card_spend)

    return savings, card, home, personal, mf, {
        "home_balance": money(home_balance),
        "pers_balance": money(pers_balance),
        "mf_units": mf_units,
        "nav": nav,
        "home_emi": home_emi,
        "pers_emi": pers_emi,
    }


# --------------------------------------------------------------------------
# Renderers - one per file format
# --------------------------------------------------------------------------

def render_xlsx(ledger: Ledger, path: Path, meta: dict) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = "Account Statement"

    bold = Font(bold=True)
    for i, (label, value) in enumerate(meta.items(), start=1):
        ws.cell(row=i, column=1, value=label).font = bold
        ws.cell(row=i, column=2, value=value)

    head_row = len(meta) + 2
    headers = ["Date", "Value Date", "Narration", "Chq/Ref No",
               "Withdrawal Amt.", "Deposit Amt.", "Closing Balance"]
    fill = PatternFill("solid", start_color="DDDDDD")
    for c, h in enumerate(headers, start=1):
        cell = ws.cell(row=head_row, column=c, value=h)
        cell.font = bold
        cell.fill = fill
        cell.alignment = Alignment(horizontal="center")

    for r, t in enumerate(ledger.txns, start=head_row + 1):
        ws.cell(row=r, column=1, value=t.when.strftime("%d/%m/%Y"))
        ws.cell(row=r, column=2, value=t.when.strftime("%d/%m/%Y"))
        ws.cell(row=r, column=3, value=t.desc)
        ws.cell(row=r, column=4, value=t.ref)
        ws.cell(row=r, column=5, value=float(t.debit) if t.debit else None)
        ws.cell(row=r, column=6, value=float(t.credit) if t.credit else None)
        ws.cell(row=r, column=7, value=float(t.balance))

    for col, width in zip("ABCDEFG", [13, 13, 58, 16, 16, 16, 18]):
        ws.column_dimensions[col].width = width

    wb.save(path)


def render_csv(ledger: Ledger, path: Path, meta: dict) -> None:
    with path.open("w", newline="", encoding="utf-8") as fh:
        w = csv.writer(fh)
        for k, v in meta.items():
            w.writerow([k, v])
        w.writerow([])
        w.writerow(["Transaction Date", "Description", "Debit", "Credit", "Balance"])
        for t in ledger.txns:
            w.writerow([
                t.when.strftime("%d-%b-%Y"), t.desc,
                f"{t.debit:.2f}" if t.debit else "",
                f"{t.credit:.2f}" if t.credit else "",
                f"{t.balance:,.2f}",
            ])


def render_pdf(ledger: Ledger, path: Path, title: str, meta: dict) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                    TableStyle)

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            leftMargin=12 * mm, rightMargin=12 * mm,
                            topMargin=14 * mm, bottomMargin=14 * mm)

    story = [Paragraph(f"<b>{title}</b>", styles["Title"]), Spacer(1, 6)]
    for k, v in meta.items():
        story.append(Paragraph(f"<b>{k}:</b> {v}", styles["Normal"]))
    story.append(Spacer(1, 10))

    data = [["Date", "Transaction Details", "Ref", "Debit", "Credit", "Balance"]]
    for t in ledger.txns:
        data.append([
            t.when.strftime("%d/%m/%Y"),
            Paragraph(t.desc, styles["BodyText"]),
            t.ref[:12],
            f"{t.debit:,.2f}" if t.debit else "",
            f"{t.credit:,.2f}" if t.credit else "",
            f"{t.balance:,.2f}",
        ])

    tbl = Table(data, colWidths=[22 * mm, 72 * mm, 22 * mm, 24 * mm, 24 * mm, 26 * mm],
                repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8e8e8")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7.5),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#999999")),
        ("ALIGN", (3, 0), (-1, -1), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]))
    story.append(tbl)
    doc.build(story)


def render_docx(ledger: Ledger, path: Path, title: str, meta: dict) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, level=1)

    for k, v in meta.items():
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(str(v))

    doc.add_paragraph()
    headers = ["Date", "Particulars", "Debit", "Credit", "Balance"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True

    for t in ledger.txns:
        row = table.add_row().cells
        row[0].text = t.when.strftime("%d-%m-%Y")
        row[1].text = t.desc
        row[2].text = f"{t.debit:,.2f}" if t.debit else ""
        row[3].text = f"{t.credit:,.2f}" if t.credit else ""
        row[4].text = f"{t.balance:,.2f}"

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    doc.save(path)


def main() -> int:
    OUT.mkdir(parents=True, exist_ok=True)
    savings, card, home, personal, mf, summary = simulate()

    s_close = savings.finalize()
    c_close = card.finalize()
    h_close = home.finalize()
    p_close = personal.finalize()
    mf.finalize()

    period = f"{START:%d-%b-%Y} to {add_months(START, MONTHS - 1).replace(day=28):%d-%b-%Y}"

    # 1. Savings account -> XLSX (the most common real-world export)
    render_xlsx(savings, OUT / "hdfc_savings_2025_2026.xlsx", {
        "Account Holder": "PANKAJ SHARMA",
        "Account Number": "50100247718842",
        "Account Type": "SAVINGS ACCOUNT - REGULAR",
        "Bank": "HDFC BANK LTD",
        "Branch": "KORAMANGALA, BENGALURU",
        "IFSC": "HDFC0000521",
        "Statement Period": period,
        "Opening Balance": float(savings.opening),
        "Closing Balance": float(s_close),
        "Currency": "INR",
    })

    # 2. Same savings data as CSV, different header names + date format
    render_csv(savings, OUT / "hdfc_savings_2025_2026.csv", {
        "Account Number": "50100247718842",
        "Statement Period": period,
        "Opening Balance": f"{savings.opening:.2f}",
        "Closing Balance": f"{s_close:.2f}",
    })

    # 3. Credit card -> PDF
    render_pdf(card, OUT / "icici_credit_card_2025_2026.pdf",
               "ICICI Bank Credit Card Statement", {
                   "Card Holder": "PANKAJ SHARMA",
                   "Card Number": "4471 XXXX XXXX 8842",
                   "Card Type": "ICICI AMAZON PAY CREDIT CARD",
                   "Statement Period": period,
                   "Credit Limit": "5,00,000.00",
                   "Opening Balance": f"{card.opening:,.2f}",
                   "Total Amount Due": f"{c_close:,.2f}",
                   "Payment Due Date": "18-Sep-2026",
               })

    # 4. Home loan -> DOCX
    render_docx(home, OUT / "hdfc_home_loan_2025_2026.docx",
                "HDFC Home Loan - Statement of Account", {
                    "Borrower Name": "PANKAJ SHARMA",
                    "Loan Account Number": "HL4471929",
                    "Loan Type": "HOME LOAN",
                    "Sanctioned Amount": "48,00,000.00",
                    "Rate of Interest": "8.75% p.a. (Floating)",
                    "EMI Amount": f"{summary['home_emi']:,.2f}",
                    "Opening Principal Outstanding": f"{home.opening:,.2f}",
                    "Closing Principal Outstanding": f"{summary['home_balance']:,.2f}",
                    "Original Tenure": "240 months",
                    "Statement Period": period,
                })

    # 5. Personal loan -> XLSX
    render_xlsx(personal, OUT / "bajaj_personal_loan_2025_2026.xlsx", {
        "Borrower Name": "PANKAJ SHARMA",
        "Loan Account Number": "PL8823641",
        "Loan Type": "PERSONAL LOAN",
        "Lender": "BAJAJ FINSERV",
        "Rate of Interest": "14.25% p.a.",
        "EMI Amount": float(summary["pers_emi"]),
        "Opening Principal Outstanding": float(personal.opening),
        "Closing Principal Outstanding": float(summary["pers_balance"]),
        "Statement Period": period,
    })

    # 6. Mutual fund -> PDF
    render_pdf(mf, OUT / "mf_portfolio_statement_2025_2026.pdf",
               "Mutual Fund Consolidated Account Statement", {
                   "Investor Name": "PANKAJ SHARMA",
                   "PAN": "ABCDE1234F",
                   "Folio Number": "91827364/55",
                   "Statement Period": period,
                   "Total Units Held": f"{summary['mf_units']:.4f}",
                   "Latest NAV": f"{summary['nav']:.4f}",
                   "Current Value": f"{summary['mf_units'] * summary['nav']:,.2f}",
               })

    print("Generated sample statements in", OUT)
    for f in sorted(OUT.iterdir()):
        print(f"  {f.name:44s} {f.stat().st_size / 1024:8.1f} KB")

    print("\nGround truth (what the pipeline must reproduce):")
    print(f"  savings   opening {savings.opening:>14,.2f}  closing {s_close:>14,.2f}  rows {len(savings.txns):>4}")
    print(f"  card      opening {card.opening:>14,.2f}  closing {c_close:>14,.2f}  rows {len(card.txns):>4}")
    print(f"  home loan opening {home.opening:>14,.2f}  closing {h_close:>14,.2f}  rows {len(home.txns):>4}")
    print(f"  personal  opening {personal.opening:>14,.2f}  closing {p_close:>14,.2f}  rows {len(personal.txns):>4}")
    print(f"  mf units  {summary['mf_units']:.4f} @ NAV {summary['nav']}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
