"""Format-specific extractors.

Each extractor has exactly one job: get raw tabular rows plus surrounding text
out of a file. None of them know what a transaction is - that is the
normalizer's job. This split is what lets a new file format be supported by
adding one function here and nothing else anywhere.

Every extractor returns an `ExtractionResult` and must never raise for a
malformed file; it reports the problem in `warnings` and returns what it got.
"""

from __future__ import annotations

import csv
import io
import logging
import re
from pathlib import Path

from ..models.schemas import ExtractedTable, ExtractionResult, SourceFormat

log = logging.getLogger(__name__)

#: Rows shorter than this are page furniture, not data.
MIN_TABLE_COLUMNS = 3
MIN_TABLE_ROWS = 2

#: How many dated rows a candidate needs before we stop looking for a better
#: table. Below this, the expensive `stream` extraction is worth running.
MIN_DATED_ROWS = 5


def _best_dated_row_count(tables: list) -> int:
    """Most dated rows found in any single candidate table.

    A cheap proxy for "did we already find the transaction list?". Uses the
    same date parser the normalizer will, so the two agree on what counts.
    """
    from ..normalize.parsers import parse_date

    best = 0
    for table in tables:
        for column in range(min(3, max((len(r) for r in table.rows), default=0))):
            hits = sum(
                1 for row in table.rows
                if column < len(row) and parse_date(row[column]) is not None
            )
            best = max(best, hits)
    return best


def _collapse_doubled(token: str) -> str:
    """Undo character doubling in a single token.

    Some issuers render bold text by drawing every glyph twice, and PDF text
    extraction faithfully returns both copies: "STATEMENT DATE" comes out as
    "SSTTAATTEEMMEENNTT DDAATTEE". Every keyword match, date parse and amount
    parse then fails on what is otherwise a perfectly good statement.

    The test is exact rather than fuzzy: a doubled token has identical
    even- and odd-indexed characters. "SSTTAA" -> "STA" == "STA". A real word
    with a double letter ("SUCCESS") fails that test, so nothing legitimate is
    corrupted.
    """
    if len(token) < 4 or len(token) % 2 != 0:
        return token
    if token[0::2] == token[1::2]:
        return token[0::2]
    return token


def collapse_doubled_text(text: str) -> str:
    """Apply doubled-glyph repair, without destroying legitimate values.

    The naive version corrupts data: "1122" is a perfectly good amount, but it
    passes the doubled test and collapses to "12". Silently turning 1122 into 12
    in a financial ledger is far worse than leaving a cosmetic artifact.

    So a purely numeric token is only repaired when the WHOLE line is doubled -
    which is what a real rendering artifact looks like. An isolated number in an
    otherwise normal line is left exactly as it is.
    """
    if not text:
        return text

    compact = text.replace(" ", "")
    # A run of pure digits is never treated as a doubled line: "11223344" is an
    # 8-digit value that would collapse to "1234". Requiring a letter or a
    # separator keeps genuine artifacts ("1122//1100//22002255") repairable
    # while leaving bare numbers untouched.
    has_non_digit = any(not c.isdigit() for c in compact)
    line_is_doubled = (
        len(compact) >= 8
        and len(compact) % 2 == 0
        and has_non_digit
        and compact[0::2] == compact[1::2]
    )
    if line_is_doubled:
        return " ".join(_collapse_doubled(t) for t in text.split(" "))

    # Otherwise repair only tokens containing letters, where a doubled run is
    # unambiguous and no numeric value is at risk.
    return " ".join(
        _collapse_doubled(t) if any(c.isalpha() for c in t) else t
        for t in text.split(" ")
    )


def _clean_cell(value) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    # PDF extraction leaves newlines inside cells where text wrapped.
    text = text.replace("\n", " ").replace("\r", " ")
    text = re.sub(r"\s+", " ", text)
    return collapse_doubled_text(text)


def _clean_rows(rows: list[list]) -> list[list[str]]:
    cleaned = []
    for row in rows:
        cells = [_clean_cell(c) for c in row]
        if any(cells):  # drop fully blank rows
            cleaned.append(cells)
    return cleaned


# --------------------------------------------------------------------------
# CSV / TSV
# --------------------------------------------------------------------------

def extract_csv(path: Path) -> ExtractionResult:
    result = ExtractionResult(extractor_used="csv", source_format=SourceFormat.CSV)

    raw: str | None = None
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            raw = path.read_text(encoding=encoding)
            break
        except (UnicodeDecodeError, LookupError):
            continue
    if raw is None:
        result.warnings.append("Could not decode file with any known encoding")
        return result

    # Sniff the delimiter; bank exports use comma, semicolon, tab or pipe.
    sample = raw[:8192]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        delimiter = dialect.delimiter
    except csv.Error:
        counts = {d: sample.count(d) for d in ",;\t|"}
        delimiter = max(counts, key=counts.get) if any(counts.values()) else ","

    rows = _clean_rows(list(csv.reader(io.StringIO(raw), delimiter=delimiter)))
    if not rows:
        result.warnings.append("File contained no rows")
        return result

    result.full_text = raw
    # A CSV is one logical table, but bank exports prepend metadata lines that
    # are narrower than the data. Split on the widest contiguous block.
    for block in _split_by_width(rows):
        if len(block) >= MIN_TABLE_ROWS and len(block[0]) >= MIN_TABLE_COLUMNS:
            result.tables.append(ExtractedTable(
                rows=block,
                confidence=0.9,
                surrounding_text=raw[:2000],
            ))
    if not result.tables:
        result.tables.append(ExtractedTable(rows=rows, confidence=0.5,
                                            surrounding_text=raw[:2000]))
    return result


def _row_span(row: list[str]) -> int:
    """Index of the last non-empty cell, plus one.

    Span is used instead of "count of filled cells" because a transaction row
    fills either the debit or the credit column, never both, so its filled-count
    oscillates while its span stays constant. Splitting on filled-count tears a
    single transaction table into fragments and silently loses rows.
    """
    last = -1
    for i, cell in enumerate(row):
        if str(cell or "").strip():
            last = i
    return last + 1


def _split_by_width(rows: list[list[str]], tolerance: int = 1) -> list[list[list[str]]]:
    """Group consecutive rows of similar span into blocks.

    Statement exports look like:  metadata (2 cols) / blank / table (7 cols).
    Splitting on span separates the letterhead block from the transaction block
    without needing to know anything about the content.
    """
    blocks: list[list[list[str]]] = []
    current: list[list[str]] = []
    prev_span = -1

    for row in rows:
        span = _row_span(row)
        if span == 0:
            continue
        # Compared against the previous row only - a running max would make one
        # wide row poison every narrower row that follows it.
        if current and abs(span - prev_span) > tolerance:
            blocks.append(current)
            current = [row]
        else:
            current.append(row)
        prev_span = span

    if current:
        blocks.append(current)

    blocks.sort(key=len, reverse=True)
    return blocks


# --------------------------------------------------------------------------
# Excel (.xlsx / .xlsm / .xls)
# --------------------------------------------------------------------------

def extract_excel(path: Path) -> ExtractionResult:
    suffix = path.suffix.lower()
    fmt = SourceFormat.XLS if suffix == ".xls" else SourceFormat.XLSX
    result = ExtractionResult(extractor_used="excel", source_format=fmt)

    if suffix == ".xls":
        return _extract_xls_legacy(path, result)

    try:
        from openpyxl import load_workbook
    except ImportError:  # pragma: no cover
        result.warnings.append("openpyxl not installed")
        return result

    try:
        # data_only=True gives us computed values instead of formula strings.
        wb = load_workbook(path, data_only=True, read_only=True)
    except Exception as exc:
        result.warnings.append(f"Could not open workbook: {exc}")
        return result

    text_parts: list[str] = []
    try:
        for sheet in wb.worksheets:
            rows = _clean_rows([
                [_excel_cell(c) for c in row]
                for row in sheet.iter_rows(values_only=True)
            ])
            if not rows:
                continue
            text_parts.append("\n".join("\t".join(r) for r in rows[:30]))

            for block in _split_by_width(rows):
                if len(block) >= MIN_TABLE_ROWS and max(len(r) for r in block) >= MIN_TABLE_COLUMNS:
                    result.tables.append(ExtractedTable(
                        rows=block,
                        source_sheet=sheet.title,
                        confidence=0.95,  # Excel cells have no layout ambiguity
                        surrounding_text="\n".join(text_parts[-1:])[:2000],
                    ))
    finally:
        wb.close()

    result.full_text = "\n\n".join(text_parts)
    if not result.tables:
        result.warnings.append("No tabular blocks found in workbook")
    return result


def _excel_cell(value) -> str:
    """Render a cell without letting Excel's float storage corrupt an amount.

    openpyxl hands back 123456.78000000001 for some currency cells. Formatting
    through repr would poison the Decimal parse, so numbers are normalized here.
    """
    if value is None:
        return ""
    if isinstance(value, bool):
        return str(value)
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return f"{value:.4f}".rstrip("0").rstrip(".")
    if hasattr(value, "strftime"):
        return value.strftime("%Y-%m-%d")
    return str(value).strip()


def _extract_xls_legacy(path: Path, result: ExtractionResult) -> ExtractionResult:
    try:
        import xlrd
    except ImportError:  # pragma: no cover
        result.warnings.append("xlrd not installed; cannot read legacy .xls")
        return result

    try:
        book = xlrd.open_workbook(str(path))
    except Exception as exc:
        result.warnings.append(f"Could not open .xls: {exc}")
        return result

    for sheet in book.sheets():
        rows = _clean_rows([
            [sheet.cell_value(r, c) for c in range(sheet.ncols)]
            for r in range(sheet.nrows)
        ])
        if not rows:
            continue
        for block in _split_by_width(rows):
            if len(block) >= MIN_TABLE_ROWS:
                result.tables.append(ExtractedTable(
                    rows=block, source_sheet=sheet.name, confidence=0.9))
    return result


# --------------------------------------------------------------------------
# Word (.docx)
# --------------------------------------------------------------------------

def extract_docx(path: Path) -> ExtractionResult:
    result = ExtractionResult(extractor_used="docx", source_format=SourceFormat.DOCX)

    try:
        from docx import Document
    except ImportError:  # pragma: no cover
        result.warnings.append("python-docx not installed")
        return result

    try:
        doc = Document(str(path))
    except Exception as exc:
        result.warnings.append(f"Could not open document: {exc}")
        return result

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    result.full_text = "\n".join(paragraphs)

    for table in doc.tables:
        rows = _clean_rows([[cell.text for cell in row.cells] for row in table.rows])
        if len(rows) >= MIN_TABLE_ROWS and max((len(r) for r in rows), default=0) >= MIN_TABLE_COLUMNS:
            result.tables.append(ExtractedTable(
                rows=rows,
                confidence=0.9,  # docx tables are structured, not inferred
                surrounding_text=result.full_text[:2000],
            ))

    # Some loan statements ship as plain paragraphs with no table at all.
    if not result.tables and paragraphs:
        parsed = _rows_from_text_lines(paragraphs)
        if parsed:
            result.tables.append(ExtractedTable(
                rows=parsed, confidence=0.4,
                surrounding_text=result.full_text[:2000]))
            result.warnings.append("No Word table found; recovered rows from paragraph text")

    if not result.tables:
        result.warnings.append("No tables or table-like text found in document")
    return result


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------

def _resolve_pdf_password(
    path: Path,
    password: str | None,
    candidates: list[str] | None,
) -> tuple[str | None, bool]:
    """Find a password that opens the PDF.

    Returns (working_password, is_encrypted). An unencrypted file returns
    (None, False). An encrypted file we cannot open returns (None, True), which
    the caller surfaces as needs_password.

    Candidates are the user's own derived passwords (see ingestion.passwords);
    they are tried only after any explicit password the caller passed.
    """
    from pypdf import PdfReader

    try:
        reader = PdfReader(str(path))
    except Exception:
        return password, False  # not our problem to diagnose here

    if not reader.is_encrypted:
        return None, False

    attempts: list[str] = []
    if password:
        attempts.append(password)
    attempts += (candidates or [])
    attempts.append("")  # some PDFs are "encrypted" with an empty owner password

    for candidate in attempts:
        try:
            if reader.decrypt(candidate):  # truthy on success
                return candidate, True
        except Exception:
            continue
    return None, True


def extract_pdf(
    path: Path,
    password: str | None = None,
    password_candidates: list[str] | None = None,
) -> ExtractionResult:
    """Extract with a strategy ladder, stopping at the first that yields tables.

    1. pdfplumber lattice  - for statements with ruled table borders
    2. pdfplumber stream   - for whitespace-aligned tables with no rules
    3. text line parsing   - last resort for pure-text statements

    The chosen strategy is recorded so a bad parse can be diagnosed later.
    """
    result = ExtractionResult(extractor_used="pdf", source_format=SourceFormat.PDF)

    try:
        import pdfplumber
    except ImportError:  # pragma: no cover
        result.warnings.append("pdfplumber not installed")
        return result

    # Resolve a working password before handing off to pdfplumber, trying the
    # user's own derived candidates for a protected file. Filename-derived
    # account fragments are appended here because the caller builds the
    # candidate list once for the whole run, while this part is per file.
    candidates = list(password_candidates or [])
    if candidates:
        from .passwords import filename_number_fragments
        extra = filename_number_fragments(path.name)
        candidates += [f for f in extra if f not in candidates]

    working_password, is_encrypted = _resolve_pdf_password(
        path, password, candidates
    )
    if is_encrypted and working_password is None:
        result.needs_password = True
        result.warnings.append(
            "PDF is password protected and none of the derived passwords opened "
            "it. Add or correct your profile details, or supply the password."
        )
        return result

    if is_encrypted:
        from .passwords import redact_candidate
        result.warnings.append(
            f"Opened a protected PDF using a derived password "
            f"({redact_candidate(working_password)})."
        )

    try:
        pdf = pdfplumber.open(str(path), password=working_password or "")
    except Exception as exc:
        message = str(exc).lower()
        if "password" in message or "encrypt" in message:
            result.needs_password = True
            result.warnings.append("PDF is password protected")
        else:
            result.warnings.append(f"Could not open PDF: {exc}")
        return result

    strategies = [
        ("lattice", {"vertical_strategy": "lines", "horizontal_strategy": "lines"}),
        ("stream", {"vertical_strategy": "text", "horizontal_strategy": "text",
                    "intersection_tolerance": 5}),
    ]

    try:
        text_parts = [
            collapse_doubled_text(page.extract_text() or "") for page in pdf.pages
        ]
        result.full_text = "\n".join(text_parts)

        # Run EVERY strategy and keep all candidates, rather than stopping at the
        # first that yields any table at all.
        #
        # Stopping early is wrong for real statements: card issuers put worked
        # fee examples ("Purchase on Sep 20, 2023 ... 2,000.00") in ruled tables
        # on their terms pages, while the actual transactions are laid out as
        # plain text. Lattice finds the examples, declares success, and the real
        # data is never looked at. The normalizer already scores candidates by
        # how many rows parse as dated transactions, so handing it everything
        # lets the genuine table win on merit.
        used: list[str] = []

        def run_strategy(name: str, settings: dict) -> None:
            for page_no, page in enumerate(pdf.pages, start=1):
                try:
                    found = page.extract_tables(table_settings=settings) or []
                except Exception as exc:  # one bad page must not kill the parse
                    log.debug("page %s failed under %s: %s", page_no, name, exc)
                    continue
                for raw_table in found:
                    rows = _clean_rows(raw_table)
                    if len(rows) < MIN_TABLE_ROWS:
                        continue
                    if max((len(r) for r in rows), default=0) < MIN_TABLE_COLUMNS:
                        continue
                    result.tables.append(ExtractedTable(
                        rows=rows,
                        source_page=page_no,
                        confidence=0.85 if name == "lattice" else 0.65,
                        surrounding_text=text_parts[page_no - 1][:2000],
                    ))
                    if name not in used:
                        used.append(name)

        # Lattice is cheap and handles ruled statements, so it always runs.
        run_strategy(*strategies[0])

        # Text-line recovery always runs too, rather than being held back as a
        # last resort. For a text-laid-out statement it is the ONLY strategy
        # that finds the transactions, and it is nearly free.
        lines_ = [ln for ln in result.full_text.split("\n") if ln.strip()]
        text_rows = _rows_from_text_lines(
            lines_, month_years=infer_month_years(result.full_text)
        )
        if text_rows:
            result.tables.append(ExtractedTable(
                rows=text_rows,
                confidence=0.55,
                surrounding_text=result.full_text[:2000],
            ))
            used.append("text-lines")

        # `stream` reconstructs cells from raw character positions and is by far
        # the slowest pass - several seconds per page on a dense statement. It
        # only earns that cost when the cheap strategies found nothing that
        # looks like dated transaction rows, so it stays conditional.
        if _best_dated_row_count(result.tables) < MIN_DATED_ROWS:
            run_strategy(*strategies[1])

        result.extractor_used = f"pdf:{'+'.join(used)}" if used else "pdf"

        if not result.tables:
            result.warnings.append(
                "No transaction table could be extracted. The PDF may be a scan "
                "requiring OCR."
            )
    finally:
        pdf.close()

    return result


#: A statement line usually starts with a date and ends with amounts.
_TEXT_LINE = re.compile(
    # The date may be followed by a column separator rather than whitespace.
    # HDFC card statements render a "DATE & TIME" column as "23/08/2025| 17:47",
    # and requiring \s+ here rejected every transaction line in the file.
    r"^\s*(?P<date>"
    r"\d{1,2}[-/.][\w]{2,9}[-/.]\d{2,4}"      # 23/08/2025, 15-Jan-2026
    r"|\d{1,2}\s+[A-Za-z]{3,9}\s+\d{2,4}"     # 21 Sep 25  (IDFC cards)
    r"|\d{4}-\d{2}-\d{2}"                     # 2025-08-23
    r"|\d{1,2}[A-Za-z]{3}(?![A-Za-z0-9])"       # 03NOV  (HSBC cards, no year)
    r")"
    # Normally a separator follows the date, but PDF text extraction sometimes
    # runs it straight into the description: "01-08-2025CMS TRANSACTION ...".
    # Requiring whitespace silently DROPPED those rows, and a dropped row shows
    # up later as an unexplained reconciliation gap. A letter immediately after
    # the date is therefore accepted too - but not a digit, which would be
    # ambiguous with a longer number.
    r"(?:[\s|,;]+|(?=[A-Za-z]))"
    r"(?P<rest>.+?)$"
)

#: An amount at the end of a line, optionally followed by a short marker glyph.
#: HDFC appends a one-letter "purchase indicator" after the amount ("C 267.11 l"),
#: which would otherwise stop the amount being recognised as trailing. Cr/Dr is
#: matched INSIDE the capture group so a direction marker is still preserved.
_TRAILING_AMOUNTS = re.compile(
    r"((?:[\d,]+\.\d{2}|\(\d[\d,]*\.\d{2}\))(?:\s*(?:Cr|Dr))?)"
    r"(?:\s+[A-Za-z]{1,2})?\s*$",
    re.IGNORECASE,
)


#: A month name paired with a 4-digit year, e.g. "24 OCT 2025", "NOV 2025".
_MONTH_YEAR = re.compile(
    r"\b([A-Za-z]{3})[A-Za-z]*[\s,-]+((?:19|20)\d{2})\b", re.IGNORECASE
)
_YEARLESS_DATE = re.compile(r"^(\d{1,2})([A-Za-z]{3})$")


def infer_month_years(text: str) -> dict[str, str]:
    """Map each month abbreviation to the year it refers to in this document.

    HSBC card statements date their rows "03NOV" with no year at all, so the
    year has to come from context. Taking a single document-wide year would be
    wrong for the statements that straddle a year boundary ("24 DEC 2025 To
    23 JAN 2026"), so each month is resolved independently from the full dates
    printed in the header.
    """
    found: dict[str, str] = {}
    for month, year in _MONTH_YEAR.findall(text or ""):
        found.setdefault(month.upper()[:3], year)
    return found


def _is_continuation(line: str) -> bool:
    """A line that carries narration only - no date of its own, no amount.

    Anything with a date starts its own row, and anything ending in an amount is
    a data row, so neither can be borrowed as another row's description.
    """
    text = (line or "").strip()
    if not text or _TEXT_LINE.match(text):
        return False
    if _TRAILING_AMOUNTS.search(text):
        return False
    # Needs some actual words to be worth attaching.
    return sum(1 for c in text if c.isalpha()) >= 4


def _wrapped_description(
    lines: list[str], index: int, span: int = 1, consumed: set[int] | None = None
) -> str:
    """Narration from the lines immediately around a bare data row.

    `consumed` records which lines have already been claimed. A continuation
    line belongs to exactly one transaction, so once a row has taken it no
    later row may take it again.
    """
    parts: list[str] = []
    offsets = [*range(index - span, index), *range(index + 1, index + 1 + span)]
    for offset in offsets:
        if not (0 <= offset < len(lines)):
            continue
        if consumed is not None and offset in consumed:
            continue
        if _is_continuation(lines[offset]):
            parts.append(lines[offset].strip())
            if consumed is not None:
                consumed.add(offset)
    return " ".join(parts)[:200]


#: A description that opens with a slash-delimited token, i.e. one that starts
#: partway through a narration rather than at its beginning.
_SLASH_NARRATION_TAIL = re.compile(r"^[A-Za-z][\w.@-]*/")


def _is_narration_fragment(description: str) -> bool:
    """Whether a row's own text is a slice from the middle of a narration.

    Two signatures, both from ICICI wrapping a long remark around its dated
    line: text cut mid-word ("...CUBYTS TECHNOLOGIES PRIVATE LIMI-") and text
    that opens partway through a slash narration ("Bank/001581210828/IBL...").
    """
    text = description.strip()
    # Short markers like "B/F" match the slash shape by accident. Treating one
    # as a fragment let the opening-balance row swallow the narration of the
    # salary credit printed directly beneath it.
    if len(text) < 8:
        return False
    return text.endswith("-") or bool(_SLASH_NARRATION_TAIL.match(text))


def _rejoin_fragment(
    lines: list[str], index: int, description: str, consumed: set[int] | None = None
) -> str:
    """Re-join a slash-delimited narration that was split across the line above.

    ICICI wraps a long remark AROUND its dated line, so the row itself carries
    only the middle of the narration:

        UPI/IndianClea/bsestarmfrzp@i/PayviaRazo/ICICI
        01-10-2025 Bank/001581210828/IBL897436...  1,30,000.00  47,043.91

    Read on its own that row says "Bank/001581210828/..." - no payee, no
    merchant - so a 1.3 lakh mutual-fund purchase (bsestarmfrzp = BSE StAR MF)
    looked like uncategorised spending, as did every CRED card-bill payment.
    This was the largest opaque bucket in the ledger at 14.2 lakh.

    Requiring the line above to be a multi-part slash narration, and this row to
    begin with another slash-delimited token, keeps the join away from layouts
    that wrap BELOW a row - there the line above belongs to the previous
    transaction and stealing it would corrupt both.
    """
    parts: list[str] = []

    above_index = index - 1
    if above_index >= 0 and not (consumed is not None and above_index in consumed):
        above = lines[above_index].strip()
        # A row cut mid-word needs only a continuation above it. One that merely
        # STARTS with a slash token is weaker evidence, so it additionally
        # requires the line above to be a multi-part slash narration - that
        # keeps the join off layouts wrapping BELOW a row, where the line above
        # belongs to the previous transaction and stealing it corrupts both.
        strict = not description.strip().endswith("-")
        if _is_continuation(above) and (not strict or above.count("/") >= 2):
            parts.append(above)
            if consumed is not None:
                consumed.add(above_index)

    parts.append(description)

    # Only text cut mid-word continues onto the line below. A row that merely
    # STARTS partway through a slash narration is already complete to its right
    # - the line beneath it opens the NEXT transaction, and taking it would
    # label this row with the following row's payee.
    if description.strip().endswith("-"):
        below_index = index + 1
        if below_index < len(lines) and not (
            consumed is not None and below_index in consumed
        ):
            below = lines[below_index].strip()
            if _is_continuation(below):
                parts.append(below)
                if consumed is not None:
                    consumed.add(below_index)

    return " ".join(parts)[:200]


def _rows_from_text_lines(
    lines: list[str], month_years: dict[str, str] | None = None
) -> list[list[str]]:
    """Recover pseudo-rows from unstructured text.

    Splits each date-led line into [date, description, ...trailing amounts].
    Deliberately conservative: a line that does not start with a date is
    skipped entirely rather than guessed at.
    """
    rows: list[list[str]] = []
    #: Continuation lines already attached to a row, so no later row re-uses one.
    consumed: set[int] = set()
    for index, line in enumerate(lines):
        m = _TEXT_LINE.match(line)
        if not m:
            continue
        rest = m.group("rest")

        amounts: list[str] = []
        # Peel amounts off the right-hand side, at most three (dr, cr, balance).
        for _ in range(3):
            am = _TRAILING_AMOUNTS.search(rest)
            if not am:
                break
            amounts.insert(0, am.group(1).strip())
            rest = rest[: am.start()].rstrip()

        if not amounts:
            continue
        # A yearless "03NOV" is completed here, where the document's own header
        # dates are still available, rather than downstream where they are not.
        raw_date = m.group("date")
        bare = _YEARLESS_DATE.match(raw_date)
        if bare:
            year = (month_years or {}).get(bare.group(2).upper())
            if not year:
                continue  # no way to date this row honestly - drop it
            raw_date = f"{bare.group(1)}{bare.group(2)}{year}"

        description = rest.strip()
        if not description:
            # Some layouts wrap the narration around the data row, leaving the
            # dated line as nothing but figures:
            #
            #     NEFT-KKBKN6...-CUBYTS TECHNOLOGIES        <- description
            #     01-09-2025  1,64,561.00  1,69,986.47      <- the data row
            #     PRIVATE LIMI-JITESHSALAUG25//CMS2-...     <- description cont.
            #
            # Dropping it loses the narration entirely, so a salary credit
            # arrives with no description, matches no rule, and never counts as
            # income. Stitch the adjacent continuation lines back on.
            description = _wrapped_description(lines, index, consumed=consumed)
        elif _is_narration_fragment(description):
            description = _rejoin_fragment(lines, index, description, consumed=consumed)

        rows.append([raw_date, description, *amounts])

    # One row is enough here, unlike a generic table. Every row that reaches
    # this point already began with a parseable date AND ended with a parseable
    # amount, so it is far more constrained than an arbitrary two-row block -
    # and a low-activity month with a single transaction is real data, not noise.
    return rows
